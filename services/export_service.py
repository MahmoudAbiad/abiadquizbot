# services/export_service.py
"""
خدمة تصدير الكويزات إلى ملفات Word (.docx) و PDF بتنسيق داخلي جميل.
- تكتشف لغة الكويز (عربي/إنكليزي) تلقائياً من نص الأسئلة.
- العربي: اتجاه الصفحة/الفقرات من اليمين لليسار (RTL) بالكامل.
- الإنكليزي: اتجاه عادي من اليسار لليمين (LTR).
- الأسئلة والاختيارات فقط داخل المتن، بدون أي تلميح/شرح.
- جدول "الإجابات الصحيحة" (يتضمن الشرح) يوضع في آخر الملف فقط.
"""
import io
import os
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from logger import get_logger, log_error
from services.image_quiz_renderer import render_question_image, looks_arabic
from services.latex_text import latex_to_plain

logger = get_logger(__name__)

# ==================== إعدادات عامة ====================

ARABIC_LETTERS = ["أ", "ب", "ج", "د", "هـ", "و", "ز", "ح"]
ENGLISH_LETTERS = ["A", "B", "C", "D", "E", "F", "G", "H"]

_ARABIC_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]")
_LETTER_RE = re.compile(r"[^\W\d_]", re.UNICODE)

FONTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "fonts")


STYLE_SIMPLE = "simple"
STYLE_MODERN = "modern"
STYLE_ACADEMIC = "academic"
DEFAULT_STYLE = STYLE_SIMPLE
VALID_STYLES = {STYLE_SIMPLE, STYLE_MODERN, STYLE_ACADEMIC}
STYLE_LABELS_AR = {
    STYLE_SIMPLE: "بسيط وأنيق",
    STYLE_MODERN: "عصري وملوّن",
    STYLE_ACADEMIC: "أكاديمي كلاسيكي",
}
# أكواد قصيرة (حرف واحد) للستايل - مستخدمة بـ callback_data تبع تيليغرام لأنو محدود بـ 64 بايت
STYLE_CODE_TO_NAME = {"s": STYLE_SIMPLE, "m": STYLE_MODERN, "a": STYLE_ACADEMIC}
STYLE_NAME_TO_CODE = {v: k for k, v in STYLE_CODE_TO_NAME.items()}


def normalize_style(style: str) -> str:
    return style if style in VALID_STYLES else DEFAULT_STYLE


class ExportError(Exception):
    """خطأ عام أثناء توليد ملف التصدير (يُعرض للمستخدم برسالة ودّية)."""


_PARENS_RE = re.compile(r"\([^()]*\)|\uFF08[^\uFF08\uFF09]*\uFF09")
_LATEX_SPAN_RE = re.compile(r"\$[^$]+\$")


def _is_math_question(q: Dict[str, Any]) -> bool:
    """يكتشف هل السؤال من نمط الكويز المصوّر LaTeX - عبر علم is_math الصريح، أو
    كخط دفاع ثانٍ لو كويز قديم محفوظ بالمفضلة قبل إضافة هذا العلم: وجود جدول بيانات
    أو مقطع $...$ فعلي بنص السؤال/الخيارات كافٍ لاعتباره سؤالاً رياضياً يحتاج رسماً
    كصورة بدل نص خام."""
    if q.get("is_math"):
        return True
    if q.get("table"):
        return True
    question = str(q.get("question", ""))
    if _LATEX_SPAN_RE.search(question):
        return True
    return any(_LATEX_SPAN_RE.search(str(opt)) for opt in (q.get("options") or []))


def _strip_parens(text: str) -> str:
    """يشيل أي نص محصور بين قوسين (زي ترجمة أو توضيح) قبل حساب لغة الكويز -
    حتى ترجمة عربية بين قوسين ضمن سؤال إنكليزي (أو العكس) ما تأثر على تحديد
    اللغة الأساسية للمستند كله."""
    prev = None
    while prev != text:
        prev = text
        text = _PARENS_RE.sub(" ", text)
    return text


def _render_math_question_image_bytes(q: Dict[str, Any], idx0: int, total: int, is_ar: bool) -> Optional[bytes]:
    """يرسم صورة السؤال الرياضي (بنفس محرك عرض الكويز الفعلي services/image_quiz_renderer.py)
    لتضمينها بملف التصدير بدل كتابة رموز LaTeX خام كنص. فشل آمن: يرجع None لو تعذّر الرسم
    لأي سبب، فيسقط المستدعي تلقائياً للمسار النصي العادي بدل تعطيل التصدير بأكمله."""
    try:
        return render_question_image(q, idx0, total, is_ar)
    except Exception as e:
        log_error(logger, f"Failed rendering math question image for export (idx={idx0}): {e}", exception=e)
        return None


def detect_language(questions: List[Dict[str, Any]]) -> str:
    """يكتشف لغة الكويز: 'ar' أو 'en' اعتماداً على نسبة الأحرف العربية في عينة من النص
    (بعد تجاهل أي نص بين قوسين، مثل الترجمات أو التوضيحات)."""
    sample_parts = []
    for q in questions[:12]:
        sample_parts.append(_strip_parens(str(q.get("question", ""))))
        sample_parts.extend(_strip_parens(str(o)) for o in (q.get("options") or []))
    sample = " ".join(sample_parts)

    letters = _LETTER_RE.findall(sample)
    if not letters:
        return "en"
    arabic_count = sum(1 for ch in letters if _ARABIC_RE.match(ch))
    return "ar" if (arabic_count / len(letters)) > 0.3 else "en"


def detect_text_language(text: str) -> str:
    """يكتشف لغة نص حر (تفريغ محاضرة/ملخص...، وليس أسئلة كويز): 'ar' أو 'en' اعتماداً
    على نسبة الأحرف العربية بعيّنة من بداية النص (بعد تجاهل أي نص بين قوسين، مثل
    ترجمات أو توضيحات مدسوسة داخل النص الأساسي)."""
    sample = _strip_parens(text[:4000])
    letters = _LETTER_RE.findall(sample)
    if not letters:
        return "en"
    arabic_count = sum(1 for ch in letters if _ARABIC_RE.match(ch))
    return "ar" if (arabic_count / len(letters)) > 0.3 else "en"


_MD_BULLET_RE = re.compile(r"^\s*[-*•]\s+(.*)$")
_MD_H1_RE = re.compile(r"^#\s+(.*)$")
_MD_H2_RE = re.compile(r"^#{2,3}\s+(.*)$")


def _parse_markdown_blocks(text: str) -> List[Tuple[str, str]]:
    """يحلّل نص Markdown بسيط (عناوين # / ## أو ### ونقاط - أو * أو •) إلى قائمة
    (نوع, محتوى) جاهزة للعرض بملفات Word/PDF. تُدمج الأسطر المتتالية غير الفارغة
    التي لا تطابق أي نمط عنوان/نقطة بنفس الفقرة، حتى أول سطر فارغ أو عنصر جديد -
    هيك ما تنكسر فقرة واحدة طويلة إلى أسطر منفصلة بالمستند النهائي."""
    blocks: List[Tuple[str, str]] = []
    para_buffer: List[str] = []

    def _flush() -> None:
        if para_buffer:
            blocks.append(("para", " ".join(para_buffer).strip()))
            para_buffer.clear()

    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            _flush()
            continue
        h2_match = _MD_H2_RE.match(line) if line.startswith("##") else None
        h1_match = _MD_H1_RE.match(line) if (not h2_match and line.startswith("#")) else None
        bullet_match = _MD_BULLET_RE.match(line)
        if h1_match:
            _flush()
            blocks.append(("h1", h1_match.group(1).strip()))
        elif h2_match:
            _flush()
            blocks.append(("h2", h2_match.group(1).strip()))
        elif bullet_match:
            _flush()
            blocks.append(("bullet", bullet_match.group(1).strip()))
        else:
            para_buffer.append(line)
    _flush()
    return blocks


def build_export_filename(title: str, ext: str) -> str:
    """اسم ملف آمن (بدون رموز قد تكسر بعض الأنظمة) مع الحفاظ على العربية."""
    safe = re.sub(r'[\\/:*?"<>|]+', " ", title or "quiz").strip()
    safe = re.sub(r"\s+", "_", safe)[:60] or "quiz"
    return f"{safe}.{ext}"


def _letters_for(is_ar: bool) -> List[str]:
    return ARABIC_LETTERS if is_ar else ENGLISH_LETTERS


# ==================== DOCX ====================

def _set_paragraph_rtl(paragraph, align_right: bool = True) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align_right else WD_ALIGN_PARAGRAPH.LEFT


def _style_run(run, font_name: str, size_pt: float, bold: bool = False,
                color: Tuple[int, int, int] = None, rtl: bool = False,
                italic: bool = False, cs_font: str = "Arial") -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    # خط الكتابة المعقّدة (w:cs) بيرسم أي حرف عربي جوا نفس الـ run، حتى لو باقي
    # الفقرة بخط لا يدعم العربي. افتراضياً Arial، بس قابل للتخصيص (مثلاً Amiri
    # لستايل "أكاديمي") حتى تنعكس هوية الخط على أي عربي مدسوس جوا نص إنكليزي.
    rFonts.set(qn("w:cs"), cs_font)
    if rtl:
        rPr.append(OxmlElement("w:rtl"))


def _set_table_rtl(table) -> None:
    tblPr = table._tbl.tblPr
    tblPr.append(OxmlElement("w:bidiVisual"))


def _set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))  # 1cm ≈ 567 twips
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_paragraph(paragraph, hex_color: str) -> None:
    """يلوّن خلفية فقرة كاملة (تمتد من هامش لهامش) - مفيدة كـ"بانر" ملوّن تحت العنوان."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _shade_run(run, hex_color: str) -> None:
    """يلوّن خلفية جزء نص واحد فقط (run) - مستخدم لعمل تأثير "بادج" حول حرف الخيار."""
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def _set_paragraph_border_bottom(paragraph, hex_color: str = "AAAAAA", size: int = 6) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_borders(cell, hex_color: str, size: int = 8) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tcPr.append(borders)


def _rgb_to_hex(rgb: Tuple[int, int, int]) -> str:
    return "%02X%02X%02X" % rgb


# إعدادات كل ستايل شكلي (ألوان/خطوط/عناصر) لملف الـ Word
_DOCX_STYLES = {
    STYLE_SIMPLE: dict(
        accent=(20, 55, 110), muted=(90, 90, 90), en_font="Calibri", ar_font="Arial",
        question_box=False, badge=False, divider=False, header_banner=False,
        table_header_bg="D9E2F3", table_header_fg=None, table_zebra=None,
    ),
    STYLE_MODERN: dict(
        accent=(37, 99, 235), muted=(100, 116, 139), en_font="Calibri", ar_font="Arial",
        question_box=True, question_box_bg="EEF4FF", question_box_border="93C5FD",
        badge=True, badge_bg=(37, 99, 235), divider=False, header_banner=True,
        table_header_bg="2563EB", table_header_fg=(255, 255, 255), table_zebra="F1F5F9",
    ),
    STYLE_ACADEMIC: dict(
        accent=(15, 23, 42), muted=(70, 70, 70), en_font="Times New Roman", ar_font="Amiri",
        question_box=False, badge=False, divider=True, header_banner=False,
        table_header_bg="E5E5E5", table_header_fg=None, table_zebra=None,
    ),
}


def _add_table_cell_text(cell, text: str, font_name: str, size_pt: float, is_ar: bool,
                          bold: bool = False, color: Tuple[int, int, int] = None,
                          cs_font: str = "Arial") -> None:
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    if is_ar:
        _set_paragraph_rtl(p, align_right=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _style_run(run, font_name, size_pt, bold=bold, color=color, rtl=is_ar, cs_font=cs_font)


def build_quiz_docx(title: str, questions: List[Dict[str, Any]], style: str = DEFAULT_STYLE) -> bytes:
    """يبني ملف Word جميل التنسيق للكويز، مع اتجاه تلقائي حسب اللغة وستايل شكلي قابل للاختيار
    (simple / modern / academic)."""
    if not questions:
        raise ExportError("لا توجد أسئلة لتصديرها.")

    style = normalize_style(style)
    cfg = _DOCX_STYLES[style]

    is_ar = detect_language(questions) == "ar"
    letters = _letters_for(is_ar)
    body_font = cfg["ar_font"] if is_ar else cfg["en_font"]
    cs_font = cfg["ar_font"]
    accent_hex = _rgb_to_hex(cfg["accent"])

    doc = Document()
    section = doc.sections[0]
    usable_width_cm = (section.page_width - section.left_margin - section.right_margin) / 360000

    # اتجاه الصفحة بالكامل RTL للعربي
    if is_ar:
        for sec in doc.sections:
            sec._sectPr.append(OxmlElement("w:bidi"))

    # ==================== العنوان ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_ar:
        _set_paragraph_rtl(title_p, align_right=False)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title or ("كويز" if is_ar else "Quiz"))
    title_color = (255, 255, 255) if cfg["header_banner"] else cfg["accent"]
    _style_run(title_run, body_font, 20, bold=True, color=title_color, rtl=is_ar, cs_font=cs_font)
    if cfg["header_banner"]:
        _shade_paragraph(title_p, accent_hex)
        title_p.paragraph_format.space_before = Pt(6)
        title_p.paragraph_format.space_after = Pt(6)

    sub_p = doc.add_paragraph()
    sub_text = f"عدد الأسئلة: {len(questions)}" if is_ar else f"Total Questions: {len(questions)}"
    sub_run = sub_p.add_run(sub_text)
    _style_run(sub_run, body_font, 11, color=cfg["muted"], rtl=is_ar, cs_font=cs_font)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_ar:
        _set_paragraph_rtl(sub_p, align_right=False)
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if style == STYLE_ACADEMIC:
        # سطر معلومات شبيه بورقة امتحان رسمية: الاسم / التاريخ
        info_p = doc.add_paragraph()
        info_text = "الاسم: ________________________     التاريخ: ______________" if is_ar \
            else "Name: ________________________     Date: ______________"
        info_run = info_p.add_run(info_text)
        _style_run(info_run, body_font, 11, color=(30, 30, 30), rtl=is_ar, cs_font=cs_font)
        if is_ar:
            _set_paragraph_rtl(info_p, align_right=True)
        info_p.paragraph_format.space_before = Pt(10)
        _set_paragraph_border_bottom(info_p, "333333", size=10)
        info_p.paragraph_format.space_after = Pt(14)
    else:
        doc.add_paragraph()

    # ==================== الأسئلة ====================
    total_questions = len(questions)
    for idx, q in enumerate(questions, 1):
        q_text = str(q.get("question", "")).strip()
        options = q.get("options") or []
        label = f"السؤال {idx}: " if is_ar else f"Question {idx}: "

        # 🆕 سؤال رياضي (نمط الكويز المصوّر LaTeX، أو جدول بيانات): يُرسم كصورة
        # نظيفة بنفس محرك عرض الكويز الفعلي بدل كتابة رموز LaTeX الخام كنص -
        # هيك تُصدَّر المسألة كاملة (نص + معادلات + جدول لو وُجد) في وحدة واحدة
        # غير قابلة للتقسيم أو التشويه، تماماً كما تظهر للطالب أثناء حل الكويز.
        if _is_math_question(q):
            img_bytes = _render_math_question_image_bytes(q, idx - 1, total_questions, is_ar)
            if img_bytes:
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                img_run.add_picture(io.BytesIO(img_bytes), width=Cm(min(usable_width_cm, 15)))
                img_p.paragraph_format.space_after = Pt(12)
                continue
            # فشل الرسم لسبب ما - نكمل للمسار النصي العادي أدناه كخط أمان أخير
            # (نظّف علامات $ ورموز LaTeX الخام أولاً بدل عرضها كما هي)
            q_text = latex_to_plain(q_text)
            options = [latex_to_plain(str(o)) for o in options]

        if cfg["question_box"]:
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            if is_ar:
                _set_table_rtl(tbl)
            cell = tbl.cell(0, 0)
            _set_cell_width(cell, usable_width_cm)
            _shade_cell(cell, cfg["question_box_bg"])
            _set_cell_borders(cell, cfg["question_box_border"])

            qp = cell.paragraphs[0]
            if is_ar:
                _set_paragraph_rtl(qp, align_right=True)
            run = qp.add_run(label + q_text)
            _style_run(run, body_font, 13, bold=True, color=cfg["accent"], rtl=is_ar, cs_font=cs_font)
            qp.paragraph_format.space_after = Pt(6)

            for opt_idx, opt_text in enumerate(options):
                letter = letters[opt_idx] if opt_idx < len(letters) else str(opt_idx + 1)
                op = cell.add_paragraph()
                if is_ar:
                    _set_paragraph_rtl(op, align_right=True)
                if cfg["badge"]:
                    badge_run = op.add_run(f" {letter} ")
                    _style_run(badge_run, body_font, 11.5, bold=True, color=(255, 255, 255), rtl=is_ar, cs_font=cs_font)
                    _shade_run(badge_run, _rgb_to_hex(cfg["badge_bg"]))
                    sep_run = op.add_run("  ")
                    _style_run(sep_run, body_font, 11.5, rtl=is_ar, cs_font=cs_font)
                    text_run = op.add_run(opt_text)
                    _style_run(text_run, body_font, 11.5, rtl=is_ar, cs_font=cs_font)
                else:
                    run = op.add_run(f"{letter}) {opt_text}")
                    _style_run(run, body_font, 11.5, rtl=is_ar, cs_font=cs_font)
                op.paragraph_format.space_after = Pt(3)

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(10)
        else:
            qp = doc.add_paragraph()
            run = qp.add_run(label + q_text)
            _style_run(run, body_font, 13, bold=True,
                       color=cfg["accent"] if style == STYLE_ACADEMIC else None, rtl=is_ar, cs_font=cs_font)
            if is_ar:
                _set_paragraph_rtl(qp, align_right=True)
            qp.paragraph_format.space_after = Pt(6)

            for opt_idx, opt_text in enumerate(options):
                letter = letters[opt_idx] if opt_idx < len(letters) else str(opt_idx + 1)
                op = doc.add_paragraph()
                sep = "." if style == STYLE_ACADEMIC else ")"
                run = op.add_run(f"{letter}{sep} {opt_text}")
                _style_run(run, body_font, 11.5, rtl=is_ar, cs_font=cs_font)
                if is_ar:
                    _set_paragraph_rtl(op, align_right=True)
                    op.paragraph_format.right_indent = Cm(0.8)
                else:
                    op.paragraph_format.left_indent = Cm(0.8)
                op.paragraph_format.space_after = Pt(2)

            spacer = doc.add_paragraph()
            if cfg["divider"]:
                spacer.paragraph_format.space_before = Pt(4)
                _set_paragraph_border_bottom(spacer, "AAAAAA", size=6)
            spacer.paragraph_format.space_after = Pt(10)

    # ==================== جدول الإجابات الصحيحة ====================
    doc.add_page_break()
    head_p = doc.add_paragraph()
    head_text = "🗝️ جدول الإجابات الصحيحة" if is_ar else "🗝️ Answer Key"
    head_run = head_p.add_run(head_text)
    _style_run(head_run, body_font, 16, bold=True, color=cfg["accent"], rtl=is_ar, cs_font=cs_font)
    if is_ar:
        _set_paragraph_rtl(head_p, align_right=True)
    head_p.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if is_ar:
        _set_table_rtl(table)

    headers = ["#", "الإجابة الصحيحة", "الشرح"] if is_ar else ["#", "Correct Answer", "Explanation"]
    widths = [1.5, 5.0, 9.5]
    hdr_cells = table.rows[0].cells
    for cell, htext, w in zip(hdr_cells, headers, widths):
        _add_table_cell_text(cell, htext, body_font, 11.5, is_ar, bold=True, color=cfg["table_header_fg"], cs_font=cs_font)
        _shade_cell(cell, cfg["table_header_bg"])
        _set_cell_width(cell, w)

    for idx, q in enumerate(questions, 1):
        options = q.get("options") or []
        correct_idx = q.get("correct_option_id")
        try:
            correct_idx = int(correct_idx)
        except (TypeError, ValueError):
            correct_idx = -1
        correct_letter = letters[correct_idx] if 0 <= correct_idx < len(letters) else "-"
        correct_text = latex_to_plain(str(options[correct_idx])) if 0 <= correct_idx < len(options) else "-"
        explanation = latex_to_plain(str(q.get("explanation") or "").strip()) or ("لا يوجد شرح" if is_ar else "No explanation")

        row_cells = table.add_row().cells
        values = [str(idx), f"{correct_letter}) {correct_text}", explanation]
        for cell, val, w in zip(row_cells, values, widths):
            _add_table_cell_text(cell, val, body_font, 10.5, is_ar, cs_font=cs_font)
            _set_cell_width(cell, w)
        if cfg["table_zebra"] and idx % 2 == 0:
            for cell in row_cells:
                _shade_cell(cell, cfg["table_zebra"])

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def build_document_docx(title: str, text: str, style: str = DEFAULT_STYLE) -> bytes:
    """يبني ملف Word منسّق لنص حر كامل (تفريغ محاضرة، ملخص، أي نص تعليمي عام) بعنوان
    وفقرات، مع اكتشاف تلقائي للغة (عربي/إنكليزي)، دعم RTL كامل للعربي (اتجاه الصفحة/
    الفقرات وخط الكتابة المعقّدة w:cs)، وطباعة عربية/إنكليزية سليمة - بنفس الستايلات
    الشكلية (simple/modern/academic) المتوفرة بتصدير الكويز. يدعم Markdown بسيط: عناوين
    (# / ## / ###) تُعرض كعناوين مُنسّقة، ونقاط (- أو * أو •) تُعرض كقائمة نقطية، وباقي
    الأسطر تُدمج كفقرات نصّية عادية."""
    if not text or not text.strip():
        raise ExportError("لا يوجد نص لتصديره.")

    style = normalize_style(style)
    cfg = _DOCX_STYLES[style]

    is_ar = detect_text_language(text) == "ar"
    body_font = cfg["ar_font"] if is_ar else cfg["en_font"]
    cs_font = cfg["ar_font"]
    accent_hex = _rgb_to_hex(cfg["accent"])

    doc = Document()

    # اتجاه الصفحة بالكامل RTL للعربي
    if is_ar:
        for sec in doc.sections:
            sec._sectPr.append(OxmlElement("w:bidi"))

    # ==================== العنوان ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_ar:
        _set_paragraph_rtl(title_p, align_right=False)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title or ("مستند" if is_ar else "Document"))
    title_color = (255, 255, 255) if cfg["header_banner"] else cfg["accent"]
    _style_run(title_run, body_font, 20, bold=True, color=title_color, rtl=is_ar, cs_font=cs_font)
    if cfg["header_banner"]:
        _shade_paragraph(title_p, accent_hex)
        title_p.paragraph_format.space_before = Pt(6)
        title_p.paragraph_format.space_after = Pt(6)

    if style == STYLE_ACADEMIC:
        # سطر معلومات شبيه بغلاف مستند أكاديمي رسمي: الاسم / التاريخ
        info_p = doc.add_paragraph()
        info_text = "الاسم: ________________________     التاريخ: ______________" if is_ar \
            else "Name: ________________________     Date: ______________"
        info_run = info_p.add_run(info_text)
        _style_run(info_run, body_font, 11, color=(30, 30, 30), rtl=is_ar, cs_font=cs_font)
        if is_ar:
            _set_paragraph_rtl(info_p, align_right=True)
        info_p.paragraph_format.space_before = Pt(10)
        _set_paragraph_border_bottom(info_p, "333333", size=10)
        info_p.paragraph_format.space_after = Pt(14)
    else:
        doc.add_paragraph()

    # ==================== الفقرات (Markdown مبسّط) ====================
    for block_type, content in _parse_markdown_blocks(text):
        if not content.strip():
            continue
        p = doc.add_paragraph()
        if is_ar:
            _set_paragraph_rtl(p, align_right=True)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if block_type == "h1":
            run = p.add_run(content)
            _style_run(run, body_font, 16, bold=True, color=cfg["accent"], rtl=is_ar, cs_font=cs_font)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            if cfg["divider"]:
                _set_paragraph_border_bottom(p, "AAAAAA", size=6)
        elif block_type == "h2":
            run = p.add_run(content)
            _style_run(run, body_font, 13.5, bold=True, color=cfg["accent"], rtl=is_ar, cs_font=cs_font)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif block_type == "bullet":
            run = p.add_run(f"• {content}")
            _style_run(run, body_font, 11.5, rtl=is_ar, cs_font=cs_font)
            if is_ar:
                p.paragraph_format.right_indent = Cm(0.8)
            else:
                p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
        else:
            run = p.add_run(content)
            _style_run(run, body_font, 11.5, rtl=is_ar, cs_font=cs_font)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.3

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ==================== PDF ====================

_FONTS_REGISTERED = False
_BIDI_AVAILABLE = True
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except ImportError:  # pragma: no cover
    _BIDI_AVAILABLE = False

from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

_FONT_FILES = {
    "Arabic": "NotoNaskhArabic-Regular.ttf",
    "Arabic-Bold": "NotoNaskhArabic-Bold.ttf",
    "Latin": "NotoSans-Regular.ttf",
    "Latin-Bold": "NotoSans-Bold.ttf",
    # اختياري: يُستخدم فقط بستايل "أكاديمي" العربي إذا موجود، وإلا بيرجع تلقائياً
    # لخط Noto Naskh Arabic العادي (مافي داعي نوقف تصدير PDF لغيابه).
    "Arabic-Academic": "Amiri-Regular.ttf",
    "Arabic-Academic-Bold": "Amiri-Bold.ttf",
}


def _register_pdf_fonts() -> None:
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    for font_name, filename in _FONT_FILES.items():
        path = os.path.join(FONTS_DIR, filename)
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, path))
            except Exception as e:
                log_error(logger, f"Failed registering font {font_name}: {e}", exception=e)
    _FONTS_REGISTERED = True


def _font_available(name: str) -> bool:
    return name in pdfmetrics.getRegisteredFontNames()


def _shape(text: str, base_ar: bool) -> str:
    """يهيئ وييعيد ترتيب النص للعرض الصحيح (bidi) - يشتغل كل ما كان في نص عربي
    بالسطر، بغض النظر عن كون الفقرة ككل عربي أو إنكليزي. base_ar يحدد اتجاه
    القراءة الأساسي للسطر (يهم بترتيب الأجزاء غير-العربية المدسوسة وسطه)."""
    if not _BIDI_AVAILABLE or not _ARABIC_RE.search(text):
        return text
    try:
        return get_display(arabic_reshaper.reshape(text), base_dir=("R" if base_ar else "L"))
    except Exception:
        return text


_SCRIPT_SPLIT_RE = re.compile(
    r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+"
    r"|[^\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]+"
)


def _split_script_runs(shaped_text: str) -> List[Tuple[str, bool]]:
    """يقسّم نص (بعد التهيئة/bidi) إلى أجزاء متتالية عربي/غير-عربي، بنفس ترتيب
    الرسم من اليسار لليمين على الصفحة - كل جزء جاهز يترسم بخطّه المناسب."""
    if not shaped_text:
        return [("", False)]
    return [(tok, bool(_ARABIC_RE.match(tok))) for tok in _SCRIPT_SPLIT_RE.findall(shaped_text)]


def _mixed_width(text: str, base_ar: bool, font_en: str, font_ar: str, size: float) -> float:
    shaped = _shape(text, base_ar)
    total = 0.0
    for seg, seg_is_ar in _split_script_runs(shaped):
        total += pdfmetrics.stringWidth(seg, font_ar if seg_is_ar else font_en, size)
    return total


def _wrap_line(text: str, base_ar: bool, font_en: str, font_ar: str, size: float,
               max_width: float) -> List[str]:
    words = text.split()
    if not words:
        return [""]
    lines: List[str] = []
    current: List[str] = []
    for w in words:
        trial = current + [w]
        width = _mixed_width(" ".join(trial), base_ar, font_en, font_ar, size)
        if width <= max_width or not current:
            current = trial
        else:
            lines.append(" ".join(current))
            current = [w]
    if current:
        lines.append(" ".join(current))
    return lines


def _draw_mixed_line(c, text: str, base_ar: bool, font_en: str, font_ar: str, size: float,
                      y: float, color: str, margin: float, page_w: float,
                      extra_indent: float = 0, center: bool = False) -> None:
    """يرسم سطر واحد ممكن يحوي عربي وإنكليزي مع بعض، كل جزء بخطّه الصحيح،
    مع محاذاة يمين/يسار/وسط محسوبة على العرض الكلي الحقيقي للسطر."""
    shaped = _shape(text, base_ar)
    runs = _split_script_runs(shaped)
    total_w = sum(pdfmetrics.stringWidth(seg, font_ar if seg_is_ar else font_en, size)
                  for seg, seg_is_ar in runs)
    if center:
        x = page_w / 2 - total_w / 2
    elif base_ar:
        x = page_w - margin - extra_indent - total_w
    else:
        x = margin + extra_indent
    c.setFillColor(HexColor(color))
    for seg, seg_is_ar in runs:
        font = font_ar if seg_is_ar else font_en
        c.setFont(font, size)
        c.drawString(x, y, seg)
        x += pdfmetrics.stringWidth(seg, font, size)


# إعدادات كل ستايل شكلي للـ PDF
_PDF_STYLES = {
    STYLE_SIMPLE: dict(
        accent="#14376E", muted="#666666", header_banner=False, box=False, divider=False,
        table_header_bg="#D9E2F3", table_header_fg="#111111", table_zebra=None, academic_fonts=False,
    ),
    STYLE_MODERN: dict(
        accent="#2563EB", muted="#64748B", header_banner=True, box=True,
        box_bg="#EEF4FF", box_border="#93C5FD", divider=False,
        table_header_bg="#2563EB", table_header_fg="#FFFFFF", table_zebra="#F1F5F9", academic_fonts=False,
    ),
    STYLE_ACADEMIC: dict(
        accent="#0F172A", muted="#464646", header_banner=False, box=False, divider=True,
        table_header_bg="#E5E5E5", table_header_fg="#111111", table_zebra=None, academic_fonts=True,
    ),
}


class _QuizPDFRenderer:
    def __init__(self, title: str, questions: List[Dict[str, Any]], is_ar: bool, style: str = DEFAULT_STYLE):
        self.title = title or ("كويز" if is_ar else "Quiz")
        self.questions = questions
        self.is_ar = is_ar
        self.letters = _letters_for(is_ar)
        self.style = normalize_style(style)
        self.cfg = _PDF_STYLES[self.style]

        self.page_w, self.page_h = A4
        self.margin = 48
        self.buf = io.BytesIO()
        self.c = canvas.Canvas(self.buf, pagesize=A4)
        self.y = self.page_h - self.margin
        self.page_num = 1

        _register_pdf_fonts()
        if self.cfg["academic_fonts"] and _font_available("Arabic-Academic"):
            # خط Amiri الكلاسيكي متوفر - نستخدمه لستايل "أكاديمي" العربي تحديداً
            self.font_ar_regular = "Arabic-Academic"
            self.font_ar_bold = "Arabic-Academic-Bold" if _font_available("Arabic-Academic-Bold") \
                else "Arabic-Academic"
        else:
            self.font_ar_regular = "Arabic" if _font_available("Arabic") else None
            self.font_ar_bold = "Arabic-Bold" if _font_available("Arabic-Bold") else None
        if self.cfg["academic_fonts"]:
            # Times-Roman/Times-Bold من الخطوط الأساسية الـ 14 المدمجة بـ reportlab -
            # موجودة دايماً بدون ما نحتاج ملف خط، وبتعطي طابع أكاديمي/رسمي كلاسيكي.
            self.font_en_regular = "Times-Roman"
            self.font_en_bold = "Times-Bold"
        else:
            self.font_en_regular = "Latin" if _font_available("Latin") else "Helvetica"
            self.font_en_bold = "Latin-Bold" if _font_available("Latin-Bold") else "Helvetica-Bold"

        if is_ar:
            self.font_regular = self.font_ar_regular
            self.font_bold = self.font_ar_bold
        else:
            self.font_regular = self.font_en_regular
            self.font_bold = self.font_en_bold

        if is_ar and (not self.font_regular or not self.font_bold):
            raise ExportError(
                "تعذّر إنشاء PDF بالعربية لعدم توفر خط عربي على الخادم حالياً. "
                "جرّب تصدير Word بدلاً من ذلك."
            )

    # ---------- أدوات رسم أساسية ----------

    def _new_page(self):
        self._draw_footer()
        self.c.showPage()
        self.page_num += 1
        self.y = self.page_h - self.margin

    def _draw_footer(self):
        self.c.setFont(self.font_regular, 9)
        self.c.setFillColor(HexColor("#888888"))
        self.c.drawCentredString(self.page_w / 2, 25, str(self.page_num))

    def _ensure_space(self, needed: float):
        if self.y - needed < self.margin + 25:
            self._new_page()

    def _draw_paragraph(self, text: str, size: float, bold: bool = False,
                         color: str = "#111111", extra_indent: float = 0, gap_after: float = 6,
                         center: bool = False, is_ar: bool = None):
        # is_ar=None يعني "استخدم اتجاه المستند العام"؛ تمرير قيمة صريحة يسمح
        # برسم فقرة واحدة بعكس اتجاه بقية المستند لو احتجنا لهيك مستقبلاً.
        direction_ar = self.is_ar if is_ar is None else is_ar
        font_en = self.font_en_bold if bold else self.font_en_regular
        font_ar = (self.font_ar_bold if bold else self.font_ar_regular) or font_en
        max_width = self.page_w - 2 * self.margin - extra_indent
        lines = _wrap_line(text, direction_ar, font_en, font_ar, size, max_width)
        for line in lines:
            self._ensure_space(size + 5)
            _draw_mixed_line(self.c, line, direction_ar, font_en, font_ar, size, self.y, color,
                              self.margin, self.page_w, extra_indent=extra_indent, center=center)
            self.y -= (size + 5)
        self.y -= gap_after

    # ---------- الأقسام ----------

    def _render_header(self):
        cfg = self.cfg
        if cfg["header_banner"]:
            band_h = 46
            self._ensure_space(band_h + 20)
            self.c.setFillColor(HexColor(cfg["accent"]))
            self.c.rect(self.margin, self.y - band_h, self.page_w - 2 * self.margin, band_h, fill=1, stroke=0)
            self.y -= (band_h / 2 - 7)
            _draw_mixed_line(self.c, self.title, self.is_ar, self.font_en_bold,
                              self.font_ar_bold or self.font_en_bold, 17, self.y, "#FFFFFF",
                              self.margin, self.page_w, center=True)
            self.y -= (band_h / 2 + 7) + 14
        else:
            self._draw_paragraph(self.title, 19, bold=True, color=cfg["accent"], gap_after=4, center=True)

        sub = f"عدد الأسئلة: {len(self.questions)}" if self.is_ar else f"Total Questions: {len(self.questions)}"
        self._draw_paragraph(sub, 10.5, color=cfg["muted"], gap_after=10 if self.style == STYLE_ACADEMIC else 16,
                              center=True)

        if self.style == STYLE_ACADEMIC:
            info = "الاسم: ________________________     التاريخ: ______________" if self.is_ar \
                else "Name: ________________________     Date: ______________"
            self._draw_paragraph(info, 11, color="#1E1E1E", gap_after=2)
            self.c.setStrokeColor(HexColor("#333333"))
            self.c.setLineWidth(1)
            self.c.line(self.margin, self.y + 8, self.page_w - self.margin, self.y + 8)
            self.y -= 12

    def _draw_math_question_image(self, q: Dict[str, Any], idx0: int) -> bool:
        """يرسم صورة السؤال الرياضي (نفس محرك الرسم الفعلي) داخل صفحة الـ PDF بدل
        كتابة رموز LaTeX خام كنص - يرجع True لو نجح الرسم فعلياً، أو False لو تعذّر
        (فيسقط المستدعي تلقائياً للمسار النصي العادي كخط أمان أخير)."""
        img_bytes = _render_math_question_image_bytes(q, idx0, len(self.questions), self.is_ar)
        if not img_bytes:
            return False
        try:
            reader = ImageReader(io.BytesIO(img_bytes))
            img_w_px, img_h_px = reader.getSize()
            max_w = self.page_w - 2 * self.margin
            display_w = min(max_w, max_w * 0.92)
            display_h = display_w * (img_h_px / img_w_px)
            self._ensure_space(display_h + 14)
            x = self.margin + (max_w - display_w) / 2
            self.c.drawImage(reader, x, self.y - display_h, width=display_w, height=display_h,
                              preserveAspectRatio=True, mask="auto")
            self.y -= (display_h + 14)
            return True
        except Exception as e:
            log_error(logger, f"Failed drawing math question image in PDF (idx={idx0}): {e}", exception=e)
            return False

    def _render_questions(self):
        cfg = self.cfg
        if cfg["box"]:
            self._render_questions_boxed()
            return
        for idx, q in enumerate(self.questions, 1):
            if _is_math_question(q) and self._draw_math_question_image(q, idx - 1):
                continue
            label = f"{'السؤال' if self.is_ar else 'Question'} {idx}: {latex_to_plain(str(q.get('question', '')).strip())}"
            self._ensure_space(30)
            label_color = cfg["accent"] if self.style == STYLE_ACADEMIC else "#111111"
            self._draw_paragraph(label, 12.5, bold=True, color=label_color, gap_after=4)
            for oidx, opt in enumerate(q.get("options") or []):
                letter = self.letters[oidx] if oidx < len(self.letters) else str(oidx + 1)
                sep = "." if self.style == STYLE_ACADEMIC else ")"
                self._draw_paragraph(f"{letter}{sep} {latex_to_plain(str(opt))}", 11, extra_indent=18, gap_after=2)
            if cfg["divider"]:
                self._ensure_space(14)
                self.c.setStrokeColor(HexColor("#AAAAAA"))
                self.c.setLineWidth(0.75)
                self.c.line(self.margin, self.y, self.page_w - self.margin, self.y)
                self.y -= 6
            self.y -= 8

    def _render_questions_boxed(self):
        """ستايل 'modern': كل سؤال جوا بطاقة (مستطيل بحواف دائرية) بخلفية فاتحة."""
        cfg = self.cfg
        pad = 12
        font_en, font_ar = self.font_en_regular, (self.font_ar_regular or self.font_en_regular)
        font_en_b, font_ar_b = self.font_en_bold, (self.font_ar_bold or self.font_en_bold)
        max_w = self.page_w - 2 * self.margin - 2 * pad

        for idx, q in enumerate(self.questions, 1):
            if _is_math_question(q) and self._draw_math_question_image(q, idx - 1):
                continue
            label = f"{'السؤال' if self.is_ar else 'Question'} {idx}: {latex_to_plain(str(q.get('question', '')).strip())}"
            options = [latex_to_plain(str(o)) for o in (q.get("options") or [])]

            q_lines = _wrap_line(label, self.is_ar, font_en_b, font_ar_b, 12.5, max_w)
            content_h = 2 * pad + len(q_lines) * (12.5 + 5) + 6
            opt_line_counts = []
            for oidx, opt in enumerate(options):
                letter = self.letters[oidx] if oidx < len(self.letters) else str(oidx + 1)
                lines = _wrap_line(f"{letter}) {opt}", self.is_ar, font_en, font_ar, 11, max_w - 10)
                opt_line_counts.append(lines)
                content_h += len(lines) * (11 + 4) + 3

            self._ensure_space(content_h + 14)
            top_y = self.y
            self.c.setFillColor(HexColor(cfg["box_bg"]))
            self.c.roundRect(self.margin, top_y - content_h, self.page_w - 2 * self.margin, content_h,
                              6, fill=1, stroke=0)
            self.c.setStrokeColor(HexColor(cfg["box_border"]))
            self.c.setLineWidth(1)
            self.c.roundRect(self.margin, top_y - content_h, self.page_w - 2 * self.margin, content_h,
                              6, fill=0, stroke=1)

            self.y = top_y - pad
            self._draw_paragraph(label, 12.5, bold=True, color=cfg["accent"], extra_indent=pad, gap_after=6)
            for oidx, opt in enumerate(options):
                letter = self.letters[oidx] if oidx < len(self.letters) else str(oidx + 1)
                self._draw_paragraph(f"{letter}) {opt}", 11, extra_indent=pad + 10, gap_after=3)
            self.y = top_y - content_h - 12

    def _render_answer_table(self):
        cfg = self.cfg
        self._new_page()
        title = "🗝️ جدول الإجابات الصحيحة" if self.is_ar else "🗝️ Answer Key"
        self._draw_paragraph(title, 15, bold=True, color=cfg["accent"], gap_after=12)

        total_w = self.page_w - 2 * self.margin
        col_ratios = [0.09, 0.31, 0.60]
        col_widths = [total_w * r for r in col_ratios]
        headers = ["#", "الإجابة الصحيحة", "الشرح"] if self.is_ar else ["#", "Correct Answer", "Explanation"]

        # للعربي: أول عمود منطقي يكون في أقصى اليمين
        if self.is_ar:
            col_x_right = [self.page_w - self.margin]
            for w in col_widths[:-1]:
                col_x_right.append(col_x_right[-1] - w)
            col_x_left = [x - w for x, w in zip(col_x_right, col_widths)]
        else:
            col_x_left = [self.margin]
            for w in col_widths[:-1]:
                col_x_left.append(col_x_left[-1] + w)
            col_x_right = [x + w for x, w in zip(col_x_left, col_widths)]

        row_font_size = 9.5
        pad = 5
        font_en = self.font_en_regular
        font_ar = self.font_ar_regular or font_en
        font_en_b = self.font_en_bold
        font_ar_b = self.font_ar_bold or font_en_b

        def draw_row(values, is_header=False, zebra=False):
            fen, far = (font_en_b, font_ar_b) if is_header else (font_en, font_ar)
            # التفاف كل خلية أولاً لتحديد ارتفاع الصف
            wrapped_cols = []
            for val, w in zip(values, col_widths):
                wrapped_cols.append(_wrap_line(str(val), self.is_ar, fen, far, row_font_size, w - 2 * pad))
            n_lines = max(len(w) for w in wrapped_cols)
            row_h = n_lines * (row_font_size + 4) + 2 * pad
            self._ensure_space(row_h)

            top_y = self.y
            if is_header:
                self.c.setFillColor(HexColor(cfg["table_header_bg"]))
                self.c.rect(self.margin, top_y - row_h, total_w, row_h, fill=1, stroke=0)
            elif zebra and cfg["table_zebra"]:
                self.c.setFillColor(HexColor(cfg["table_zebra"]))
                self.c.rect(self.margin, top_y - row_h, total_w, row_h, fill=1, stroke=0)

            self.c.setStrokeColor(HexColor("#BBBBBB"))
            self.c.rect(self.margin, top_y - row_h, total_w, row_h, fill=0, stroke=1)
            for x_left in col_x_left[1:]:
                self.c.line(x_left, top_y, x_left, top_y - row_h)

            text_color = cfg["table_header_fg"] if is_header else "#111111"
            for lines, x_left, x_right in zip(wrapped_cols, col_x_left, col_x_right):
                line_y = top_y - pad - row_font_size
                for line in lines:
                    shaped = _shape(line, self.is_ar)
                    runs = _split_script_runs(shaped)
                    run_w = sum(pdfmetrics.stringWidth(seg, far if a else fen, row_font_size)
                                for seg, a in runs)
                    cx = (x_right - pad - run_w) if self.is_ar else (x_left + pad)
                    self.c.setFillColor(HexColor(text_color))
                    for seg, seg_is_ar in runs:
                        f = far if seg_is_ar else fen
                        self.c.setFont(f, row_font_size)
                        self.c.drawString(cx, line_y, seg)
                        cx += pdfmetrics.stringWidth(seg, f, row_font_size)
                    line_y -= (row_font_size + 4)
            self.y = top_y - row_h

        draw_row(headers, is_header=True)
        for idx, q in enumerate(self.questions, 1):
            options = q.get("options") or []
            correct_idx = q.get("correct_option_id")
            try:
                correct_idx = int(correct_idx)
            except (TypeError, ValueError):
                correct_idx = -1
            correct_letter = self.letters[correct_idx] if 0 <= correct_idx < len(self.letters) else "-"
            correct_text = latex_to_plain(str(options[correct_idx])) if 0 <= correct_idx < len(options) else "-"
            explanation = latex_to_plain(str(q.get("explanation") or "").strip()) or ("لا يوجد شرح" if self.is_ar else "No explanation")
            draw_row([str(idx), f"{correct_letter}) {correct_text}", explanation], zebra=(idx % 2 == 0))

    def render(self) -> bytes:
        self._render_header()
        self._render_questions()
        self._render_answer_table()
        self._draw_footer()
        self.c.save()
        self.buf.seek(0)
        return self.buf.getvalue()


def build_quiz_pdf(title: str, questions: List[Dict[str, Any]], style: str = DEFAULT_STYLE) -> bytes:
    """يبني ملف PDF جميل التنسيق للكويز، مع اتجاه تلقائي حسب اللغة وستايل شكلي قابل للاختيار
    (يتطلب خطوط assets/fonts للعربي)."""
    if not questions:
        raise ExportError("لا توجد أسئلة لتصديرها.")
    is_ar = detect_language(questions) == "ar"
    renderer = _QuizPDFRenderer(title, questions, is_ar, style=style)
    return renderer.render()


class _DocumentPDFRenderer:
    """يبني PDF لنص حر (تفريغ محاضرة، ملخص، أي نص تعليمي عام) بترقيم صفحات وتشكيل
    bidi صحيح للعربي، باستخدام نفس منطق تسجيل الخطوط (Noto/Amiri) وتقسيم النص إلى
    أجزاء عربي/إنكليزي المستخدم بمُصدِّر الكويز - مكرَّر هون عمداً بدل الوراثة من
    _QuizPDFRenderer حتى يبقى مسار الكويز الحالي دون أي تعديل."""

    def __init__(self, title: str, text: str, is_ar: bool, style: str = DEFAULT_STYLE):
        self.title = title or ("مستند" if is_ar else "Document")
        self.blocks = _parse_markdown_blocks(text)
        self.is_ar = is_ar
        self.style = normalize_style(style)
        self.cfg = _PDF_STYLES[self.style]

        self.page_w, self.page_h = A4
        self.margin = 48
        self.buf = io.BytesIO()
        self.c = canvas.Canvas(self.buf, pagesize=A4)
        self.y = self.page_h - self.margin
        self.page_num = 1

        _register_pdf_fonts()
        if self.cfg["academic_fonts"] and _font_available("Arabic-Academic"):
            # خط Amiri الكلاسيكي متوفر - نستخدمه لستايل "أكاديمي" العربي تحديداً
            self.font_ar_regular = "Arabic-Academic"
            self.font_ar_bold = "Arabic-Academic-Bold" if _font_available("Arabic-Academic-Bold") \
                else "Arabic-Academic"
        else:
            self.font_ar_regular = "Arabic" if _font_available("Arabic") else None
            self.font_ar_bold = "Arabic-Bold" if _font_available("Arabic-Bold") else None
        if self.cfg["academic_fonts"]:
            # Times-Roman/Times-Bold من الخطوط الأساسية الـ 14 المدمجة بـ reportlab -
            # موجودة دايماً بدون ما نحتاج ملف خط، وبتعطي طابع أكاديمي/رسمي كلاسيكي.
            self.font_en_regular = "Times-Roman"
            self.font_en_bold = "Times-Bold"
        else:
            self.font_en_regular = "Latin" if _font_available("Latin") else "Helvetica"
            self.font_en_bold = "Latin-Bold" if _font_available("Latin-Bold") else "Helvetica-Bold"

        if is_ar:
            self.font_regular = self.font_ar_regular
            self.font_bold = self.font_ar_bold
        else:
            self.font_regular = self.font_en_regular
            self.font_bold = self.font_en_bold

        if is_ar and (not self.font_regular or not self.font_bold):
            raise ExportError(
                "تعذّر إنشاء PDF بالعربية لعدم توفر خط عربي على الخادم حالياً. "
                "جرّب تصدير Word بدلاً من ذلك."
            )

    # ---------- أدوات رسم أساسية (نفس منطق _QuizPDFRenderer) ----------

    def _new_page(self):
        self._draw_footer()
        self.c.showPage()
        self.page_num += 1
        self.y = self.page_h - self.margin

    def _draw_footer(self):
        self.c.setFont(self.font_regular, 9)
        self.c.setFillColor(HexColor("#888888"))
        self.c.drawCentredString(self.page_w / 2, 25, str(self.page_num))

    def _ensure_space(self, needed: float):
        if self.y - needed < self.margin + 25:
            self._new_page()

    def _draw_paragraph(self, text: str, size: float, bold: bool = False,
                         color: str = "#111111", extra_indent: float = 0, gap_after: float = 6,
                         center: bool = False):
        font_en = self.font_en_bold if bold else self.font_en_regular
        font_ar = (self.font_ar_bold if bold else self.font_ar_regular) or font_en
        max_width = self.page_w - 2 * self.margin - extra_indent
        lines = _wrap_line(text, self.is_ar, font_en, font_ar, size, max_width)
        for line in lines:
            self._ensure_space(size + 5)
            _draw_mixed_line(self.c, line, self.is_ar, font_en, font_ar, size, self.y, color,
                              self.margin, self.page_w, extra_indent=extra_indent, center=center)
            self.y -= (size + 5)
        self.y -= gap_after

    # ---------- الأقسام ----------

    def _render_header(self):
        cfg = self.cfg
        if cfg["header_banner"]:
            band_h = 46
            self._ensure_space(band_h + 20)
            self.c.setFillColor(HexColor(cfg["accent"]))
            self.c.rect(self.margin, self.y - band_h, self.page_w - 2 * self.margin, band_h, fill=1, stroke=0)
            self.y -= (band_h / 2 - 7)
            _draw_mixed_line(self.c, self.title, self.is_ar, self.font_en_bold,
                              self.font_ar_bold or self.font_en_bold, 17, self.y, "#FFFFFF",
                              self.margin, self.page_w, center=True)
            self.y -= (band_h / 2 + 7) + 14
        else:
            self._draw_paragraph(self.title, 19, bold=True, color=cfg["accent"], gap_after=14, center=True)

        if self.style == STYLE_ACADEMIC:
            info = "الاسم: ________________________     التاريخ: ______________" if self.is_ar \
                else "Name: ________________________     Date: ______________"
            self._draw_paragraph(info, 11, color="#1E1E1E", gap_after=2)
            self.c.setStrokeColor(HexColor("#333333"))
            self.c.setLineWidth(1)
            self.c.line(self.margin, self.y + 8, self.page_w - self.margin, self.y + 8)
            self.y -= 12

    def _render_body(self):
        cfg = self.cfg
        for block_type, content in self.blocks:
            if not content.strip():
                continue
            if block_type == "h1":
                self._ensure_space(30)
                self._draw_paragraph(content, 15, bold=True, color=cfg["accent"], gap_after=8)
                if cfg["divider"]:
                    self.c.setStrokeColor(HexColor("#AAAAAA"))
                    self.c.setLineWidth(0.75)
                    self.c.line(self.margin, self.y + 4, self.page_w - self.margin, self.y + 4)
                    self.y -= 4
            elif block_type == "h2":
                self._ensure_space(24)
                self._draw_paragraph(content, 13, bold=True, color=cfg["accent"], gap_after=6)
            elif block_type == "bullet":
                self._draw_paragraph(f"• {content}", 11, color="#111111",
                                      extra_indent=16, gap_after=4)
            else:
                self._draw_paragraph(content, 11, color="#111111", gap_after=8)

    def render(self) -> bytes:
        self._render_header()
        self._render_body()
        self._draw_footer()
        self.c.save()
        self.buf.seek(0)
        return self.buf.getvalue()


def build_document_pdf(title: str, text: str, style: str = DEFAULT_STYLE) -> bytes:
    """يبني ملف PDF منسّق لنص حر كامل (تفريغ محاضرة، ملخص...) بخطوط Noto Naskh Arabic/
    Amiri للعربي وNoto Sans/Times للإنكليزي، تشكيل bidi صحيح لأي نص عربي مدسوس وسط
    نص إنكليزي أو العكس، وترقيم صفحات تلقائي - بنفس الستايلات الشكلية المتوفرة
    بتصدير الكويز (simple / modern / academic)."""
    if not text or not text.strip():
        raise ExportError("لا يوجد نص لتصديره.")
    is_ar = detect_text_language(text) == "ar"
    renderer = _DocumentPDFRenderer(title, text, is_ar, style=style)
    return renderer.render()
